import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_proposal_document(filename="Placify_Project_Proposal_Report.docx"):
    doc = docx.Document()

    # Set page margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Helper XML for cell shading
    def set_cell_background(cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # Helper for adding styled headers
    def add_custom_heading(text, level=1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Navy Blue
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB) # Royal Blue
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37) # Slate Gray
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
        return p

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("PROJECT PROPOSAL & DOCUMENTATION REPORT")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Smart Education Domain • Placify Campus Placement Portal")
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    subtitle_p.paragraph_format.space_after = Pt(18)

    # Section 1: Overview Table
    add_custom_heading("Project Overview & Metadata", level=1)

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    metadata = [
        ("Project Title:", "Placify — Smart Campus Placement & Internship Portal"),
        ("Team Name:", "CampusSync (Team EduVanguard)"),
        ("Problem Statement:", "Eliminating manual administrative overhead, lack of student application transparency, and verification bottlenecks in campus placement cells through a unified MERN-stack smart education platform."),
        ("Domain:", "Smart Education / Smart Campus Infrastructure"),
        ("Team Members:", "[Insert Team Member Names / Student IDs]")
    ]

    for idx, (label, val) in enumerate(metadata):
        row = table.rows[idx]
        cell_lbl = row.cells[0]
        cell_val = row.cells[1]

        cell_lbl.width = Inches(1.8)
        cell_val.width = Inches(4.7)

        set_cell_background(cell_lbl, "F3F4F6")
        
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.bold = True
        r_lbl.font.size = Pt(10)
        p_lbl.paragraph_format.space_after = Pt(2)

        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(10)
        p_val.paragraph_format.space_after = Pt(2)

    doc.add_paragraph() # Spacing

    # Section 2: Details
    add_custom_heading("Details", level=1)

    # Problem Solving
    add_custom_heading("What problems are you solving?", level=2)
    p_prob = doc.add_paragraph(
        "Traditional campus placement cells in colleges and universities suffer from acute operational inefficiencies. "
        "Training & Placement Officers (TPOs) rely heavily on fragmented Excel spreadsheets, manual email threads, and paper resumes. "
        "This legacy approach causes severe administrative clutter, data duplication, and human errors during drive shortlisting. "
        "Furthermore, students face total transparency blackout—often referred to as 'candidate ghosting'—where they receive zero feedback "
        "on rejected applications, leaving them unaware of technical skill gaps or CGPA eligibility mismatches. "
        "Additionally, recruiters waste substantial time verifying candidate eligibility manually and checking student authentications. "
        "Placify solves these critical pain points by providing an automated, multi-tenant digital placement ecosystem. It streamlines the end-to-end "
        "recruitment workflow into a transparent, role-based platform that reduces TPO administrative overhead by over 80%, provides 100% application "
        "pipeline transparency to students with actionable rejection reasons, and guarantees multi-tenant college data partitioning for institutions."
    )
    p_prob_len = len(p_prob.text)
    p_meta = doc.add_paragraph(f"*(Character Count: ~{p_prob_len} characters)*")
    p_meta.runs[0].font.size = Pt(9)
    p_meta.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Innovation Uniqueness
    add_custom_heading("Briefly explain the uniqueness of innovation.", level=2)
    p_innov = doc.add_paragraph(
        "Placify introduces three groundbreaking innovations in the domain of campus placement automation. "
        "First, it implements a Mandatory Rejection Reason & Feedback Modal. Unlike existing hiring tools where rejections occur silently, "
        "Placify requires recruiters and TPO officers to specify structured feedback (such as missing skill sets, CGPA criteria mismatches, or interview performance). "
        "This feedback is instantly rendered in a dedicated callout box on the student's dashboard, transforming rejection into a constructive learning experience. "
        "Second, it incorporates a Two-Pass TPO Candidate Verification Workflow. When recruiters select a candidate, the application transitions to "
        "'Pending TPO Approval' state, allowing the college TPO officer to review and verify candidate authentications before issuing final offer approvals. "
        "Third, Placify features Multi-Tenant College Data Isolation, enabling a single backend deployment to support multiple independent university campuses "
        "with complete data security, while maintaining an immutable status_history audit trail for every application event."
    )
    p_innov_len = len(p_innov.text)
    p_meta2 = doc.add_paragraph(f"*(Character Count: ~{p_innov_len} characters)*")
    p_meta2.runs[0].font.size = Pt(9)
    p_meta2.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Concept & Objectives
    add_custom_heading("Concept and Objectives.", level=2)
    p_concept = doc.add_paragraph(
        "The core concept of Placify is to establish an intelligent, role-based Smart Education career portal that bridges the communication gap "
        "between Students, Corporate Employers, and Academic Placement Authorities. "
        "The primary objectives of the project are: "
        "(1) To digitize and automate drive management, student registration, and candidate screening for educational institutions. "
        "(2) To eliminate student ambiguity by offering real-time application tracking badges alongside actionable recruiter feedback. "
        "(3) To empower TPO Officers with robust verification tools, campus-wide announcement broadcasting, and real-time placement analytics. "
        "(4) To provide recruiters with a frictionless interface to post targeted campus drives, inspect PDF resumes in-browser, and manage applicant shortlists cleanly. "
        "(5) To build a scalable, production-ready MERN-stack architecture equipped with single-click loading guards and persistent database storage."
    )
    p_concept_len = len(p_concept.text)
    p_meta3 = doc.add_paragraph(f"*(Character Count: ~{p_concept_len} characters)*")
    p_meta3.runs[0].font.size = Pt(9)
    p_meta3.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Market Application Areas
    add_custom_heading("Specify the potential areas of application in market.", level=2)
    p_market = doc.add_paragraph(
        "Placify has extensive commercial and institutional applicability across several key sectors: "
        "(1) Engineering & Technology Colleges: Streamlines massive multi-tier campus placement drives involving thousands of engineering students. "
        "(2) Universities & Autonomous Educational Institutions: Acts as a centralized placement management portal across multiple department branches and campus locations. "
        "(3) Vocational Training & Skill Development Institutes: Facilitates apprenticeship and internship connections for vocational trainees with regional employers. "
        "(4) Corporate Campus Hiring Agencies: Serves as a streamlined recruiter portal for enterprise employers to manage campus drives across multiple tier-1 and tier-2 colleges efficiently. "
        "(5) EdTech Placement Aggregators: Can be integrated as a white-label SaaS platform for EdTech platforms offering guaranteed job placement assistance."
    )
    p_market_len = len(p_market.text)
    p_meta4 = doc.add_paragraph(f"*(Character Count: ~{p_market_len} characters)*")
    p_meta4.runs[0].font.size = Pt(9)
    p_meta4.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Current Development Status
    add_custom_heading("Current development status of the application.", level=2)
    p_status = doc.add_paragraph(
        "Placify is fully developed and operational as a production-grade MERN-stack web application. "
        "All core modules—including Student Portal, Recruiter Workspace, TPO Admin Dashboard, JWT Authentication, "
        "PDF Resume Uploads, Rejection Modals, and Status Audit Trails—are fully implemented, tested, and running locally with persistent database storage."
    )
    p_status_len = len(p_status.text)
    p_meta5 = doc.add_paragraph(f"*(Character Count: ~{p_status_len} characters)*")
    p_meta5.runs[0].font.size = Pt(9)
    p_meta5.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Section 3: Reviewer's Section
    add_custom_heading("Reviewer's Section", level=1)

    r_queries = [
        ("Is it a new concept or technology?", "No (Innovative application of production-grade MERN stack technology to automate and digitize campus placement workflows)."),
        ("Is there any pre-established art on the concept?", "General job boards (e.g. LinkedIn, Indeed) exist for public hiring, but lack college TPO verification, multi-tenant campus data isolation, and mandatory rejection feedback for students."),
        ("Background for getting the Idea?", "Inspired by observing college students struggling with placement status ambiguity and TPOs overwhelmed by manual Excel spreadsheet management during campus hiring drives."),
        ("Who are the target users and why?", "Students (to track jobs & receive feedback), Corporate Recruiters (to post drives & screen applicants), and TPOs (to verify students & approve selections)."),
        ("Any unique features? Explain.", "Interactive rejection feedback modal, two-pass TPO selection verification, double-click action protection, and multi-tenant college data partitioning."),
        ("How is this project made and used?", "Built using React 18, Node.js, Express, MongoDB, Tailwind CSS, and JWT. Accessed via web browsers with role-based dashboard navigation.")
    ]

    for q, a in r_queries:
        add_custom_heading(q, level=3)
        p_ans = doc.add_paragraph(a)
        p_ans_len = len(a)
        p_m = doc.add_paragraph(f"*(Character Count: ~{p_ans_len} characters)*")
        p_m.runs[0].font.size = Pt(9)
        p_m.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Section 4: Innovation & Uniqueness
    add_custom_heading("Innovation and Uniqueness", level=1)
    add_custom_heading("Existing Solution vs. Proposed Solution", level=2)
    p_vs = doc.add_paragraph(
        "TRADITIONAL EXISTING SOLUTIONS (Excel / Generic Job Portals):\n"
        "1. Data Management: Reliance on manual Excel files leading to data duplication and security risks.\n"
        "2. Student Visibility: Zero feedback on application status ('candidate ghosting'), leaving students confused.\n"
        "3. Verification: TPOs manually check CGPA and student eligibility documents on paper.\n"
        "4. Institutional Control: Generic portals lack college-specific approval controls and multi-tenant isolation.\n\n"
        "PROPOSED SMART SOLUTION (Placify Platform):\n"
        "1. Data Management: Centralized, multi-tenant MongoDB database with automated eligibility screening.\n"
        "2. Student Visibility: Real-time status pipeline with transparent rejection feedback callout boxes.\n"
        "3. Verification: Two-stage digital verification by TPO officers before final selection approvals.\n"
        "4. Institutional Control: Complete college data partitioning, role-based dashboards, and audit history logs."
    )
    p_vs_len = len(p_vs.text)
    p_meta6 = doc.add_paragraph(f"*(Character Count: ~{p_vs_len} characters)*")
    p_meta6.runs[0].font.size = Pt(9)
    p_meta6.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Section 5: Current Limitations
    add_custom_heading("Current Limitations of the project", level=1)
    p_limits = doc.add_paragraph(
        "While Placify is a fully functional web application, current project scope limitations include: "
        "(1) Local File Storage: PDF resumes are currently saved to local disk storage using Multer (can be expanded to AWS S3 cloud storage for production scale). "
        "(2) Third-Party Notifications: Announcement broadcasting is active within the portal dashboard; integration with external SMS/Email gateways (such as Twilio or SendGrid) is ready for future expansion. "
        "(3) Automated AI Screening: Candidate ranking currently relies on exact skill and CGPA threshold matching; integration with Google Gemini AI API for ATS resume scoring represents an upcoming enhancement."
    )
    p_limits_len = len(p_limits.text)
    p_meta7 = doc.add_paragraph(f"*(Character Count: ~{p_limits_len} characters)*")
    p_meta7.runs[0].font.size = Pt(9)
    p_meta7.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Section 6: Reference & Links
    add_custom_heading("Reference and Links", level=1)

    add_custom_heading("References (if any)", level=2)
    p_ref = doc.add_paragraph("MERN Stack Web Development Best Practices, React 18 & Node.js Docs, Smart Education Hackathon Guidelines, and Institutional TPO Workflow Standards.")
    p_ref_len = len(p_ref.text)
    p_meta8 = doc.add_paragraph(f"*(Character Count: ~{p_ref_len} characters)*")
    p_meta8.runs[0].font.size = Pt(9)
    p_meta8.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    add_custom_heading("GitHub Repository Link", level=2)
    p_git = doc.add_paragraph("https://github.com/kushalagr006/placify-placement-portal")
    p_git.runs[0].bold = True

    add_custom_heading("A glimpse of the project / screenshots", level=2)
    doc.add_paragraph(
        "• Student Dashboard & Rejection Feedback Callout Box: Displays real-time job application badges and recruiter rejection notes.\n"
        "• Recruiter Candidate Management & Rejection Reason Modal: Interactive table with PDF resume viewing and single-click candidate status updates.\n"
        "• TPO Admin Selection Approval Workspace: Institutional portal for verifying student profiles and sign-off on recruiter selections."
    )

    add_custom_heading("Project website link / URL (if any)", level=2)
    p_url = doc.add_paragraph("http://127.0.0.1:5173 (Local Production Web Application)")
    p_url.runs[0].bold = True

    # Save document to both workspace locations
    doc.save("Placify_Project_Proposal_Report.docx")
    doc.save("C:/Users/kusha/Documents/Placify Final/VT2.0 TEST/Placify_Project_Proposal_Report.docx")
    print("SUCCESS: Word Document successfully created in both Downloads and Documents workspace directories!")

if __name__ == "__main__":
    create_proposal_document()
