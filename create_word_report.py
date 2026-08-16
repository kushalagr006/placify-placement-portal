import sys
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_report():
    doc = Document()

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style definitions
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # -------------------------------------------------------------
    # COVER PAGE / TITLE
    # -------------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("\n\nVOCATIONAL TRAINING / PROJECT REPORT\n")
    run_title.bold = True
    run_title.font.size = Pt(24)
    run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Navy Blue

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("COLLEGE PLACEMENT & INTERNSHIP PORTAL\n\n")
    run_sub.bold = True
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run("Submitted in partial fulfillment of the requirements\nfor the award of the degree of\nBachelor of Technology in Computer Science & Engineering\n\n\n\n")
    run_meta.font.size = Pt(12)
    run_meta.italic = True

    p_by = doc.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_by = p_by.add_run("Submitted By:\n[Student Name]\nRoll No: [Your Roll Number]\n\nUnder Guidance Of:\n[Guide Name / Designation]\n\nDepartment of Computer Science & Engineering\n2025 - 2026\n")
    run_by.bold = True
    run_by.font.size = Pt(12)

    doc.add_page_break()

    # -------------------------------------------------------------
    # TABLE OF CONTENTS
    # -------------------------------------------------------------
    h_toc = doc.add_heading("TABLE OF CONTENTS", level=1)
    h_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table_toc = doc.add_table(rows=1, cols=3)
    table_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_toc.rows[0].cells
    hdr_cells[0].text = "Chapter"
    hdr_cells[1].text = "Title"
    hdr_cells[2].text = "Page No."

    for cell in hdr_cells:
        set_cell_background(cell, "1E3A8A")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    toc_data = [
        ("I", "Introduction", "1"),
        ("", "1.1 Project Background & Motivation", "1"),
        ("", "1.2 Problem Statement", "2"),
        ("", "1.3 Project Objectives & System Scope", "3"),
        ("II", "Hardware and Software Requirements", "4"),
        ("", "2.1 Hardware Requirements", "4"),
        ("", "2.2 Software Requirements & Tech Stack", "5"),
        ("III", "Flow Chart / E-R Diagrams / Block Diagram", "6"),
        ("", "3.1 System Architecture & Data Flow Diagrams (Block Diagram)", "6"),
        ("", "3.2 Entity-Relationship (E-R) Diagrams & Database Design", "8"),
        ("IV", "Results & Discussions", "10"),
        ("", "4.1 System Implementation & Portal Module Features", "10"),
        ("", "4.2 API Endpoint Verification & Performance Discussion", "12"),
        ("V", "Conclusion & Scope of further work", "13"),
        ("", "5.1 Conclusion", "13"),
        ("", "5.2 Scope of Further Work", "14"),
        ("", "References", "15"),
        ("", "Acknowledgement of Vocational Training", "16"),
        ("", "(Student Copy) Offer Letter", "17"),
        ("", "Certificate", "18"),
    ]

    for ch, title, pg in toc_data:
        row_cells = table_toc.add_row().cells
        row_cells[0].text = ch
        row_cells[1].text = title
        row_cells[2].text = pg

        # Alignment
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if ch != "":
            set_cell_background(row_cells[0], "F3F4F6")
            set_cell_background(row_cells[1], "F3F4F6")
            set_cell_background(row_cells[2], "F3F4F6")
            for c in row_cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.bold = True

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER I: INTRODUCTION
    # -------------------------------------------------------------
    doc.add_heading("CHAPTER I: INTRODUCTION", level=1)

    doc.add_heading("1.1 Project Background & Motivation", level=2)
    doc.add_paragraph(
        "In higher educational institutions, campus placement drives and internship coordination represent critical operational activities. Traditionally, Training and Placement Offices (TPOs) manage placement workflows through manual processes involving physical forms, spreadsheets, and decentralized email communications.\n\n"
        "The College Placement & Internship Portal is an enterprise-grade full-stack web application designed to digitize, streamline, and automate the entire campus recruitment lifecycle. By offering dedicated role-based portals for Students, Employers (Companies), and TPO Administrators, the platform facilitates seamless job postings, application submissions, resume verification, and candidate status tracking in real time."
    )

    doc.add_heading("1.2 Problem Statement", level=2)
    doc.add_paragraph(
        "Existing manual and semi-automated placement management systems suffer from several operational bottlenecks:\n"
        "1. Communication Lag: Scattered WhatsApp messages and emails lead to missed deadlines for campus drives.\n"
        "2. Duplicate Applications: Lack of automated constraints allows candidates to re-apply repeatedly for closed drives.\n"
        "3. Tracking & Analytics: Difficulty for TPO officers to quickly measure placement percentages, average packages, and company verification status.\n"
        "4. Security & Validation Concerns: Risk of unverified grade entries (CGPAs) or unauthorized access to sensitive student resumes."
    )

    doc.add_heading("1.3 Project Objectives & System Scope", level=2)
    doc.add_paragraph(
        "The main objectives of this project are:\n"
        "• Centralized Authentication: Implement stateless JWT (JSON Web Tokens) with Bcrypt password hashing across three distinct user roles: student, company, and admin.\n"
        "• Automated Validation: Enforce database-level and application-level integrity rules (e.g., UNIQUE(student_id, job_id), CGPA boundaries).\n"
        "• Asynchronous High-Performance API: Leverage FastAPI (Python) and SQLAlchemy ORM to serve scalable REST API endpoints.\n"
        "• Modern Responsive Frontend: Build an interactive, glassmorphic single-page application using React 18 and Vite."
    )

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER II: HARDWARE AND SOFTWARE REQUIREMENTS
    # -------------------------------------------------------------
    doc.add_heading("CHAPTER II: HARDWARE AND SOFTWARE REQUIREMENTS", level=1)

    doc.add_heading("2.1 Hardware Requirements", level=2)
    p_hw = doc.add_paragraph()
    p_hw.add_run("Client-Side Requirements:\n").bold = True
    p_hw.add_run("• Processor: Intel Core i3 (10th Gen) / AMD Ryzen 3 or equivalent\n• RAM: Minimum 4 GB (8 GB recommended)\n• Storage: 500 MB free space for browser caching\n• Display: 1024x768 screen resolution minimum\n\n")
    p_hw.add_run("Server-Side Requirements:\n").bold = True
    p_hw.add_run("• Processor: Dual-Core 2.0 GHz CPU minimum\n• RAM: Minimum 4 GB RAM\n• Storage: 10 GB SSD for database files and media assets\n• Network: Broadband connection with static IP / domain")

    doc.add_heading("2.2 Software Requirements & Tech Stack", level=2)
    p_sw = doc.add_paragraph()
    p_sw.add_run("Backend Stack:\n").bold = True
    p_sw.add_run("• Python 3.11+, FastAPI v0.109+, Uvicorn ASGI Server, SQLAlchemy ORM v2.0+, Alembic, JWT (python-jose) & Passlib (Bcrypt)\n\n")
    p_sw.add_run("Frontend Stack:\n").bold = True
    p_sw.add_run("• React 18, Vite v5, Axios HTTP Client, Lucide-React Icons, Tailwind CSS\n\n")
    p_sw.add_run("Database Stack:\n").bold = True
    p_sw.add_run("• MySQL Server v8.0 / PostgreSQL v15 / SQLite 3")

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER III: FLOW CHART / E-R DIAGRAMS / BLOCK DIAGRAM
    # -------------------------------------------------------------
    doc.add_heading("CHAPTER III: FLOW CHART / E-R DIAGRAMS / BLOCK DIAGRAM", level=1)

    doc.add_heading("3.1 System Architecture & Data Flow Diagrams (Block Diagram)", level=2)
    doc.add_paragraph(
        "The system follows a modern decoupled 3-Tier Architecture:\n\n"
        "[Client Layer: React + Vite Single Page Application]\n"
        "         │\n"
        "         │  JSON Requests (HTTP Bearer JWT Token)\n"
        "         ▼\n"
        "[Application Layer: FastAPI Asynchronous Server]\n"
        "         │\n"
        "         │  SQLAlchemy ORM Mapping\n"
        "         ▼\n"
        "[Database Layer: MySQL / PostgreSQL Relational Database]"
    )

    doc.add_heading("3.2 Entity-Relationship (E-R) Diagrams & Database Design", level=2)
    doc.add_paragraph(
        "The relational database schema consists of 6 core tables:\n\n"
        "1. USERS: Central authentication repository (id, name, email, password, role).\n"
        "2. STUDENTS: Student profile attributes (student_id, user_id, branch, semester, cgpa, skills, resume, phone).\n"
        "3. COMPANIES: Employer details (company_id, user_id, company_name, hr_name, website, location, is_verified).\n"
        "4. JOBS: Campus placement postings (job_id, company_id, title, package, location, deadline, status).\n"
        "5. APPLICATIONS: Submissions link (application_id, student_id, job_id, applied_date, status).\n"
        "6. ANNOUNCEMENTS: TPO notice updates (announcement_id, title, description, created_by, created_at)."
    )

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER IV: RESULTS & DISCUSSIONS
    # -------------------------------------------------------------
    doc.add_heading("CHAPTER IV: RESULTS & DISCUSSIONS", level=1)

    doc.add_heading("4.1 System Implementation & Portal Module Features", level=2)
    doc.add_paragraph(
        "The application was implemented and tested across three dedicated user portals:\n"
        "• Student Module: Candidate login, job browsing, single-click application submission, and status tracking.\n"
        "• Employer Module: Drive posting, applicant evaluation, resume viewing, and candidate shortlisting/selection.\n"
        "• TPO Admin Module: Real-time metric analytics, registered company verification, and announcement broadcasting."
    )

    doc.add_heading("4.2 API Endpoint Verification & Performance Discussion", level=2)
    doc.add_paragraph(
        "All API routes were verified using FastAPI's interactive Swagger UI (http://localhost:8000/docs):\n"
        "• POST /auth/signup -> 201 Created\n"
        "• POST /auth/login -> 200 OK (JWT Access Token)\n"
        "• GET /student/jobs -> 200 OK\n"
        "• POST /student/jobs/{id}/apply -> 201 Created\n"
        "• POST /company/jobs -> 201 Created\n"
        "• PUT /company/applications/{id}/status -> 200 OK\n"
        "• POST /admin/announcements -> 201 Created"
    )

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER V: CONCLUSION & SCOPE OF FURTHER WORK
    # -------------------------------------------------------------
    doc.add_heading("CHAPTER V: CONCLUSION & SCOPE OF FURTHER WORK", level=1)

    doc.add_heading("5.1 Conclusion", level=2)
    doc.add_paragraph(
        "The College Placement & Internship Portal successfully addresses the inefficiencies of traditional manual placement coordination. By decoupling the architecture into an asynchronous FastAPI backend and a responsive React frontend, the system delivers high performance, secure role-based authentication, and immediate real-time tracking for students, companies, and placement officers."
    )

    doc.add_heading("5.2 Scope of Further Work", level=2)
    doc.add_paragraph(
        "Future enhancements planned include:\n"
        "1. Automated Email/SMS Gateways for instant interview scheduling alerts.\n"
        "2. AI-Powered Resume Scoring matching candidate skills to job descriptions.\n"
        "3. Automated Calendar Integration for candidate interview slot selection."
    )

    doc.add_page_break()

    # -------------------------------------------------------------
    # REFERENCES & FORMAL ATTACHMENTS
    # -------------------------------------------------------------
    doc.add_heading("REFERENCES", level=1)
    doc.add_paragraph(
        "1. FastAPI Official Documentation (https://fastapi.tiangolo.com/)\n"
        "2. React 18 & Vite Documentation (https://react.dev/)\n"
        "3. SQLAlchemy 2.0 Unified Manual (https://docs.sqlalchemy.org/)\n"
        "4. JSON Web Token (JWT) RFC 7519 Specification\n"
        "5. MySQL 8.0 Reference Manual"
    )

    doc.add_page_break()

    doc.add_heading("ACKNOWLEDGEMENT OF VOCATIONAL TRAINING", level=1)
    doc.add_paragraph(
        "I express my sincere gratitude to the Department of Computer Science & Engineering and the Training & Placement Office (TPO) for granting me the opportunity to undergo vocational training and undertake the College Placement & Internship Portal project.\n\n"
        "I extend my deep appreciation to my project guide, [Guide Name / Designation], for their invaluable guidance, technical insights, and constant encouragement throughout the development of this full-stack application.\n\n"
        "Date: [Insert Date]\n"
        "Place: [Insert Location]\n\n"
        "______________________\n"
        "[Student Name]\n"
        "Roll No: [Your Roll Number]\n"
        "Department of CSE"
    )

    doc.add_page_break()

    doc.add_heading("(STUDENT COPY) OFFER LETTER", level=1)
    doc.add_paragraph(
        "COLLEGE PLACEMENT & INTERNSHIP PORTAL\n"
        "TRAINING & PLACEMENT OFFICE\n\n"
        "Date: 01st August 2026\n\n"
        "To,\n"
        "Student Name: [Student Name]\n"
        "Roll Number: [Roll Number]\n"
        "Branch: Computer Science & Engineering\n\n"
        "SUBJECT: PROVISIONAL OFFER LETTER FOR VOCATIONAL INTERNSHIP\n\n"
        "Dear [Student Name],\n\n"
        "We are pleased to inform you that based on your technical evaluations, you have been selected for the Vocational Training Program in Full-Stack Web Development at the Placement Portal Cell.\n\n"
        "Project Assigned: College Placement & Internship Portal (FastAPI + React + MySQL)\n"
        "Duration: 6 Weeks\n\n"
        "Sincerely,\n\n"
        "______________________\n"
        "Head, Training & Placement Office"
    )

    doc.add_page_break()

    doc.add_heading("CERTIFICATE", level=1)
    doc.add_paragraph(
        "DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING\n"
        "CERTIFICATE OF COMPLETION\n\n"
        "This is to certify that Mr./Ms. [STUDENT NAME], bearing Roll No. [ROLL NUMBER], a student of B.Tech (Computer Science & Engineering), has successfully completed their Vocational Training Project entitled:\n\n"
        "\"COLLEGE PLACEMENT & INTERNSHIP PORTAL\"\n\n"
        "under the guidance of [GUIDE NAME] during the academic session 2025-2026 in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology.\n\n\n\n"
        "________________________                        ________________________\n"
        "Internal Examiner                               Head of Department (CSE)"
    )

    # Save Document
    downloads_path = os.path.join(os.path.expanduser("~"), "Downloads", "College_Placement_Portal_Project_Report.docx")
    doc.save(downloads_path)
    print(f"SUCCESSFULLY_CREATED: {downloads_path}")

if __name__ == "__main__":
    create_report()
